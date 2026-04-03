# ============================
# DOCX GENERATION UTILITY
# ============================

from docx import Document


def generate_docx(template_content, context):
    content = template_content

    for key, value in context.items():
        content = content.replace(f"{{{{{key}}}}}", str(value))

    doc = Document()
    for line in content.split("\n"):
        doc.add_paragraph(line)

    return doc


# ============================
# YOUTUBE URL UTILITIES
# ============================

import re
from urllib.parse import urlparse, parse_qs


def extract_youtube_video_id(url):
    patterns = [
        r"shorts\/([a-zA-Z0-9_-]{11})",
        r"watch\?v=([a-zA-Z0-9_-]{11})",
        r"youtu\.be\/([a-zA-Z0-9_-]{11})",
    ]

    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)

    return None


def get_youtube_embed_data(url):
    video_id = extract_youtube_video_id(url)
    if not video_id:
        return None

    return {
        "embed_url": f"https://www.youtube-nocookie.com/embed/{video_id}",
        "watch_url": f"https://www.youtube.com/watch?v={video_id}"
    }


# ============================
# DOCUMENT TEXT EXTRACTION
# ============================

import pdfplumber
from docx import Document


def extract_text_from_document(file_path):
    if file_path.endswith('.pdf'):
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
        return text.strip()

    elif file_path.endswith('.docx'):
        doc = Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs).strip()

    elif file_path.endswith('.txt'):
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read().strip()

    return ""


# =====================================================
# GEMINI CONFIGURATION (REQUIRED)
# =====================================================

import os
import google.generativeai as genai
from django.conf import settings

# Get API keys safely
GOOGLE_API_KEYS = getattr(settings, "GEMINI_API_KEYS", [])
GOOGLE_API_KEY = getattr(settings, "GEMINI_API_KEY", None)

if GOOGLE_API_KEY:
    genai.configure(api_key=GOOGLE_API_KEY)
elif GOOGLE_API_KEYS:
    genai.configure(api_key=GOOGLE_API_KEYS[0])


# =====================================================
# LEGAL DOCUMENT ANALYSIS (GEMINI)
# =====================================================

def analyze_legal_document_with_gemini(document_text):
    """
    Analyze uploaded legal document and explain in simple language.
    Output is HTML-safe.
    """

    if not document_text or len(document_text.strip()) < 100:
        return "<p><em>The document does not contain enough readable text.</em></p>"

    prompt = f"""
You are an expert Indian Legal Assistant.

TASK:
Explain the following legal document in SIMPLE language for a common person.

WHAT TO DO:
- Identify what type of document it is
- Explain purpose of the document
- Explain key clauses in bullet points
- Mention obligations and risks
- Mention termination or penalties if present

STRICT RULES:
- Do NOT give legal advice
- Do NOT suggest illegal actions
- Use ONLY these HTML tags:
<p>, <strong>, <ul>, <li>, <em>

LEGAL DOCUMENT TEXT:
{document_text}
"""

    import hashlib
    from .models import AIDocumentCache

    # Generate hash for the document text
    text_hash = hashlib.sha256(document_text.encode('utf-8', errors='ignore')).hexdigest()

    # Check Cache first
    cached = AIDocumentCache.objects.filter(text_hash=text_hash).first()
    if cached:
        print("Returning cached document analysis.")
        return cached.summary

    api_keys = getattr(settings, "GEMINI_API_KEYS", [GOOGLE_API_KEY] if GOOGLE_API_KEY else [])
    
    if not api_keys:
        return "<p><strong>AI service configuration error.</strong></p>"

    # Try with Gemini (3-key failover + wait & retry queue)
    max_retries = 2
    for attempt in range(max_retries):
        for idx, key in enumerate(api_keys):
            if not key:
                continue
                
            try:
                genai.configure(api_key=key)
                # Using gemini-2.0-flash as primary (fallback to latest if needed)
                try:
                    model = genai.GenerativeModel("gemini-2.0-flash")
                    response = model.generate_content(prompt)
                except Exception as model_err:
                    print(f"Model gemini-2.0-flash failed: {model_err}. Trying gemini-flash-latest...")
                    model = genai.GenerativeModel("gemini-flash-latest")
                    response = model.generate_content(prompt)
                
                if response and response.text:
                    # Save to Cache
                    AIDocumentCache.objects.create(text_hash=text_hash, summary=response.text)
                    return response.text
            except Exception as e:
                print(f"Gemini Doc Analysis Error (Key {idx+1}, Attempt {attempt+1}):", e)
                if "429" in str(e):
                    if idx < len(api_keys) - 1:
                        continue
                    elif attempt < max_retries - 1:
                        import time
                        time.sleep(10)
                        break
                    else:
                        break # Fall through to local fallback below
                
                # For other errors, try next key
                if idx < len(api_keys) - 1:
                    continue
                if attempt == max_retries - 1:
                    break # Fall through to local fallback

    # -- LOCAL DATABASE FALLBACK --
    from django.db.models import Q
    from .models import Law

    # Clean punctuation
    query_text = ''.join(e for e in document_text[:2000] if e.isalnum() or e.isspace())
    query_words = query_text.lower().split()
    
    stop_words = {'the', 'and', 'with', 'law', 'is', 'for', 'was', 'this', 'that'}
    keywords = [w for w in query_words if w not in stop_words and len(w) > 3]
    
    # Keyword Expansion Map
    expansion_map = {
        'accident': ['motor', 'traffic', 'vehicle', 'driving', 'collision', 'road'],
        'theft': ['stealing', 'robbery', 'burglary', 'property'],
        'murder': ['homicide', 'killing', 'death', 'body'],
        'divorce': ['spouse', 'marriage', 'separation', '13b'],
        'maintenance': ['spouse', 'allowance', 'money', 'claim']
    }
    
    search_keywords = list(set(keywords))
    for kw in keywords:
        if kw in expansion_map:
            search_keywords.extend(expansion_map[kw])

    if search_keywords:
        search_filter = Q()
        for k in search_keywords[:10]: # Limit keywords for performance
            search_filter |= Q(title__icontains=k) | Q(description__icontains=k) | Q(category__icontains=k)
        
        matches = Law.objects.filter(search_filter).distinct()[:5]
        if matches:
            response_text = '<p><strong>(Live Analysis Offline - Using Local Database Matches)</strong></p>'
            response_text += 'Based on your document, here are the most relevant laws from our database:<ul>'
            for law in matches:
                punishment = f' (Penalty: {law.punishment})' if law.punishment else ''
                response_text += f'<li style="margin-bottom:12px"><strong>{law.title}</strong>: {law.description[:400]}...{punishment}</li>'
            response_text += '</ul>'
            return response_text

    return "<p><strong>Analysis currently unavailable. Please check back shortly or see local Laws tab.</strong></p>"
