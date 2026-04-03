from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.db import IntegrityError
from django.http import HttpResponse, JsonResponse, FileResponse, HttpResponseBadRequest
from django.views.decorators.csrf import csrf_exempt
from .models import UserData, Law ,LegalTemplate, Advocate, Payment, Feedback, Awareness_content, legal_office, Appointment, SOS
from django.contrib.auth.hashers import make_password, check_password
from .forms import DynamicTemplateForm
from .utils import generate_docx
from docx import Document
from docx.shared import Inches
from django.contrib.auth.decorators import login_required

from django.utils.html import strip_tags
from .models import Law, Awareness_content
from .utils import get_youtube_embed_data



import razorpay
from django.conf import settings

razorpay_client = razorpay.Client(
    auth=(settings.RAZORPAY_KEY_ID, settings.RAZORPAY_KEY_SECRET)
)

import re
import json
import google.generativeai as genai












# admin
ADMIN_USERNAME = "admin"
ADMIN_PASSWORD = "admin"
ADMIN_USER_ID = 1

def myadmin(request):
    if request.method == 'POST':
        username = request.POST['admin_username']
        password = request.POST['admin_password']

        if username == ADMIN_USERNAME and password == ADMIN_PASSWORD:
            request.session['admin_logged_in'] = True
            request.session['admin_user_id'] = ADMIN_USER_ID
            request.session['admin_username'] = ADMIN_USERNAME
            return redirect('admin_home')
        else:
            return HttpResponse("<h3 style='color:red;text-align:center;'>Invalid admin credentials</h3>")
    return render(request, 'admin/myadmin.html')

def admin_home(request):
    if request.session.get('admin_logged_in'):

        # 🧾 Fetch all users (if needed in template)
        users = UserData.objects.all()

        # 📊 Dashboard Statistics
        total_users = UserData.objects.count()
        total_advocates = Advocate.objects.count()
        pending_advocates = Advocate.objects.filter(is_verified=False).count()
        verified_advocates = Advocate.objects.filter(is_verified=True).count()
        total_templates = LegalTemplate.objects.count()
        total_laws = Law.objects.count()
        total_transactions = Payment.objects.count()
        total_awareness = Awareness_content.objects.count()

        return render(request, 'admin/admin_home.html', {
            'admin_username': request.session.get('admin_username'),
            'admin_id': request.session.get('admin_user_id'),

            # List of users (you already use this)
            'users': users,

            # 📊 Dashboard Counts
            'total_users': total_users,
            'total_advocates': total_advocates,
            'pending_advocates': pending_advocates,
            'verified_advocates': verified_advocates,
            'total_templates': total_templates,
            'total_laws': total_laws,
            'total_transactions': total_transactions,
            'total_awareness': total_awareness,
        })
    else:
        return HttpResponse("Unauthorized", status=403)

    
def admin_logout(request):
    # Clear admin session
    request.session.flush()
    return redirect('myadmin')

#def addlaw(request):
    return render(request, 'admin/addlaw.html')







def send_feedback(request, id):
    advocate = get_object_or_404(Advocate, advocate_id=id)

    # ----------------------------------------------------
    # ✅ CHECK IF USER IS LOGGED IN USING YOUR SESSION KEY
    # ----------------------------------------------------
    if not request.session.get("user_id"):
        messages.error(request, "⚠️ You need to login first!")
        return redirect('login')

    # ----------------------------------------------------
    # ✅ USER IS LOGGED IN → SAVE FEEDBACK
    # ----------------------------------------------------
    if request.method == "POST":
        rating = request.POST.get("rating")
        message = request.POST.get("message")

        Feedback.objects.create(
            user_id=request.session["user_id"],  # <-- IMPORTANT (your session)
            advocate=advocate,
            rating=rating,
            message=message,
        )

        messages.success(request, "✅ Feedback sent successfully!")
        return redirect("advocate_detail", id=advocate.advocate_id)


    return redirect("advocate_detail", id=advocate.advocate_id)




        

            





# 🏠 Home Page
from .models import Law

def home(request):
    # ✅ keep your custom session data with a DIFFERENT name
    session_user = request.session.get('user')

    category_list = (
        Law.objects
        .values('category')
        .distinct()
    )

    contents = Awareness_content.objects.all().order_by('-upload_date')

    context = {
        'session_user': session_user,   # optional, if you need fullname
        'category_list': category_list,
        'contents': contents,
    }

    return render(request, 'home.html', context)





def view_categories(request):
    category_list = (
        Law.objects
        .values('category')
        .distinct()
    )
    return render(request, 'categories.html', {
        'category_list': category_list
    })

def adv_view_categories(request):
    category_list = (
        Law.objects
        .exclude(category__isnull=True)
        .exclude(category__exact='')
        .values('category')
        .distinct()
    )
    return render(request, 'adv_view_categories.html', {
        'category_list': category_list
    })



def advocate_home(request):
    advocate = request.session.get('advocate')
    return render(request, 'advocate_home.html', {'advocate': advocate}) if advocate else render(request, 'advocate_home.html')


# 🧾 Register Page
def register(request):
    if request.method == 'POST':
        fullname = request.POST.get('fullname')
        email = request.POST.get('email')
        language = request.POST.get('language')
        location = request.POST.get('location')
        contact_no = request.POST.get('contact_no')
        password = request.POST.get('password')

        # -------------------------------
        # PHONE NUMBER VALIDATION (ADDED)
        # -------------------------------
        if not contact_no.isdigit() or len(contact_no) != 10:
            messages.warning(request, "⚠️ Phone number must be exactly 10 digits!")
            return render(request, 'register.html')

        # -------------------------------
        # CHECK EMAIL EXIST
        # -------------------------------
        if UserData.objects.filter(email=email).exists():
            messages.error(request, "⚠️ Email already exists. Please try another one.")
            return render(request, 'register.html')

        # -------------------------------
        # CREATE USER
        # -------------------------------
        hashed_pw = make_password(password)
        try:
            UserData.objects.create(
                fullname=fullname,
                email=email,
                language=language,
                location=location,
                contact_no=contact_no,
                password=hashed_pw
            )
        except IntegrityError:
            messages.error(request, "⚠️ Email already exists. Please try another one.")
            return render(request, 'register.html')

        # -------------------------------
        # SUCCESS MESSAGE POPUP
        # -------------------------------
        messages.success(request, "✅ Registration successful!")
        return redirect('login')

    return render(request, 'register.html')



# 🔐 Login Page
from django.shortcuts import render, redirect
from django.http import HttpResponse
from django.contrib.auth.hashers import check_password
from .models import UserData


from django.shortcuts import render, redirect
from django.http import HttpResponse
from django.contrib.auth import authenticate, login
from django.contrib.auth.models import User
from django.contrib.auth.hashers import check_password
from .models import UserData


def login_view(request):
    if request.method == 'POST':
        email = request.POST.get('email')
        password = request.POST.get('password')

        try:
            # Get your custom user
            userdata = UserData.objects.get(email=email)

            # Validate password
            if check_password(password, userdata.password):

                # 🔐 Get or create Django auth user
                django_user, created = User.objects.get_or_create(
                    username=userdata.email,
                    defaults={
                        'email': userdata.email,
                        'first_name': userdata.fullname
                    }
                )

                # Ensure password sync (important)
                if created or not django_user.check_password(password):
                    django_user.set_password(password)
                    django_user.save()

                # ✅ Django login (THIS FIXES NAVBAR)
                login(request, django_user)

                # ✅ Your existing sessions (kept)
                request.session['user'] = userdata.fullname
                request.session['user_id'] = userdata.id
                request.session['user_email'] = userdata.email

                messages.success(request, "✅ Login successful!")
                return redirect('home')

            else:
                messages.error(request, "❌ Incorrect password.")
                return redirect('login')

        except UserData.DoesNotExist:
            messages.error(request, "❌ Email not found. Please register.")
            return redirect('register')

    return render(request, 'login.html')




# 🚪 Logout
def logout_view(request):
    request.session.flush()  # ✅ Clears all session data
    messages.success(request, "👋 Logged out successfully.")
    return redirect('login')


# 👤 Profile Page
def profile_view(request):
    user_fullname = request.session.get('user')
    user_id = request.session.get('user_id')

    if not user_fullname or not user_id:
        messages.warning(request, "⚠️ Please login to view your profile.")
        return redirect('login')

    user = UserData.objects.get(id=user_id)

    # --- Update Profile ---
    if 'update_profile' in request.POST:
        user.fullname = request.POST.get('fullname')
        user.email = request.POST.get('email')
        user.language = request.POST.get('language')
        user.location = request.POST.get('location')
        user.contact_no = request.POST.get('contact_no')
        user.save()
        request.session['user'] = user.fullname
        messages.success(request, "✅ Profile updated successfully!")
        return redirect('profile')

    # --- Change Password ---
    if 'change_password' in request.POST:
        old_pw = request.POST.get('old_password')
        new_pw = request.POST.get('new_password')
        confirm_pw = request.POST.get('confirm_password')

        if not check_password(old_pw, user.password):
            messages.error(request, "❌ Old password is incorrect.")
        elif new_pw != confirm_pw:
            messages.warning(request, "⚠️ New passwords do not match.")
        else:
            user.password = make_password(new_pw)
            user.save()
            messages.success(request, "✅ Password changed successfully!")

        return redirect('profile')

    # --- Fetch User's Given Reviews ---
    feedbacks = Feedback.objects.filter(user=user).order_by('-created_at')

    return render(request, 'profile.html', {'user': user, 'feedbacks': feedbacks})



# ✅ View all categories
def view_categories(request):
    categories = Law.objects.values_list('category', flat=True).distinct()

    icon_map = {
        "Human Body": "flaticon-jury",
        "Inchoate": "flaticon-lawyer-1",
        "General": "flaticon-folder",
        "State": "flaticon-libra",
        "Public Order": "flaticon-libra",
        "Economic": "flaticon-money-bag",
        "Government Asset": "flaticon-money-bag",
        "Society": "flaticon-libra",
        "Morality": "flaticon-libra",
        "Religion": "flaticon-libra",
        "Property": "flaticon-act",
        "Financial": "flaticon-jury",
        "Document": "flaticon-folder",
        "Personal Relationship": "flaticon-jury",
        "Reputation": "flaticon-jury",
    }

    category_list = []

    for c in categories:
        icon = "flaticon-folder"  # default icon for unknown categories

        for key, value in icon_map.items():
            if key.lower() in c.lower():
                icon = value
                break

        category_list.append({
            "name": c,
            "icon": icon
        })

    return render(request, "view_categories.html", {
        "category_list": category_list
    })




# ✅ View laws under selected category
def view_laws_by_category(request, category):
    laws = Law.objects.filter(category=category)
    return render(request, 'view_laws_by_category.html', {
        'laws': laws,
        'category': category
    })

def adv_view_laws_by_category(request, category):
    laws = Law.objects.filter(category=category)
    return render(request, 'adv_view_laws_by_category.html', {
        'laws': laws,
        'category': category
    })


# 🧾 Add Law Page
def addlaw(request):
    if request.method == 'POST':
        law_id = request.POST.get('law_id')
        category = request.POST.get('category')
        section_no = request.POST.get('section_no')
        title = request.POST.get('title')
        description = request.POST.get('description')
        sourcelink = request.POST.get('sourcelink')

        if Law.objects.filter(section_no=section_no).exists():
            messages.warning(request, "⚠️ Law already exists. Please update.")
            return redirect('addlaw')

        
        Law.objects.create(
            law_id=law_id,
            category=category,
            section_no=section_no,
            title=title,
            description=description,
            sourcelink=sourcelink
        )

        messages.success(request, "✅ Law added successfully!")
        return redirect('addlaw')

    return render(request, 'admin/addlaw.html')



def admin_viewlaw(request):
    laws = Law.objects.all()
    return render(request, 'admin/admin_viewlaw.html', {'laws': laws})


def admin_law_detailed_view(request, law_id):
    
        law = Law.objects.get(law_id=law_id)
        return render(request, 'admin/admin_law_detailed_view.html', {'law': law})



def admin_user_profile(request):
    if request.session.get('admin_logged_in'):
        # 🧾 Fetch all registered users
        users = UserData.objects.all()

        return render(request, 'admin/admin_user_profile.html', {
            'admin_username': request.session.get('admin_username'),
            'admin_id': request.session.get('admin_user_id'),
            'users': users,                # 👈 Pass users list to template
        })
    else:
        return HttpResponse("Unauthorized", status=403)
    



def law_delete(request, law_id):
    law = Law.objects.get(law_id=law_id)
    law.delete()

    messages.success(request, "🗑️ Law deleted successfully!")
    return redirect('admin_viewlaw')

    

    



def law_update(request, law_id):

    # ✅ Correct PK field
    law = Law.objects.get(law_id=law_id)

    if request.method == "POST" and 'law_update' in request.POST:

        law.category = request.POST.get('category')
        law.section_no = request.POST.get('section_no')
        law.title = request.POST.get('title')
        law.description = request.POST.get('description')
        law.sourcelink = request.POST.get('sourcelink') or None
        law.save()

        messages.success(request, "✅ Law updated successfully!")
        return redirect('admin_law_detailed_view', law_id=law_id)

    return render(request, 'admin/admin_law_detailed_view.html', {'law': law})




# === Legal Template Admin Views ===

def admin_template_list(request):
    # Simple session-based admin check (reuse your existing admin session)
    # if not request.session.get("admin_id"):
    #     return redirect("admin_login")
    templates = LegalTemplate.objects.all().order_by('-created_at')
    return render(request, "admin/admin_template_list.html", {"templates": templates})

def addtemplate(request):
    if request.method == "POST":

        template_name = request.POST.get("template_name")
        category = request.POST.get("category")
        description = request.POST.get("description")
        template_content = request.POST.get("template_content")
        uploaded_file = request.FILES.get("uploaded_file")  # ✅ NEW

        # ✅ Duplicate check (same template name)
        if LegalTemplate.objects.filter(template_name=template_name).exists():
            messages.warning(request, "⚠️ Template already exists. Please update.")
            return redirect('addtemplate')

        # ✅ Create new template with file
        LegalTemplate.objects.create(
            template_name=template_name,
            category=category,
            description=description,
            template_content=template_content,
            uploaded_file=uploaded_file   # ✅ NEW
        )

        # ✅ Success Popup
        messages.success(request, "✅ Template added successfully!")
        return redirect('addtemplate')

    return render(request, "admin/addtemplate.html")



# def edit_template(request, template_id):
#     # if not request.session.get("admin_id"):
#     #     return redirect("admin_login")
#     t = get_object_or_404(LegalTemplate, template_id=template_id)
#     if request.method == "POST":
#         t.template_name = request.POST.get("template_name", t.template_name).strip()
#         t.category = request.POST.get("category", t.category).strip()
#         t.description = request.POST.get("description", t.description).strip()
#         t.template_content = request.POST.get("template_content", t.template_content).strip()
#         t.save()
#         return redirect("admin_template_list")
#     return render(request, "admin/edit_template.html", {"t": t})

# def delete_template(request, template_id):
#     # if not request.session.get("admin_id"):
#     #     return redirect("admin_login")
#     t = get_object_or_404(LegalTemplate, template_id=template_id)
#     if request.method == "POST":
#         t.delete()
#         return redirect("admin_template_list")
#     return render(request, "admin/delete_template_confirm.html", {"t": t})





def template_list(request):
    templates = LegalTemplate.objects.all()
    return render(request, "templist.html", {"templates": templates})

def adv_template_list(request):
    templates = LegalTemplate.objects.all()
    return render(request, "adv_templist.html", {"templates": templates})


def fill_template(request, template_id):
    template = get_object_or_404(LegalTemplate, template_id=template_id)

    if request.method == "POST":
        form = DynamicTemplateForm(template.template_content, request.POST)

        if form.is_valid():

            # ✔ Download only the uploaded file
            if template.uploaded_file:
                file_path = template.uploaded_file.path
                file_name = template.uploaded_file.name.split('/')[-1]

                # Read and return the file as download
                with open(file_path, 'rb') as f:
                    response = HttpResponse(
                        f.read(),
                        content_type="application/octet-stream"
                    )
                    response['Content-Disposition'] = f'attachment; filename="%s"' % file_name
                    return response

            else:
                return HttpResponse("No file uploaded by admin.")

    else:
        form = DynamicTemplateForm(template.template_content)

    return render(request, "filltemplate.html", {
        "form": form,
        "template": template
    })


def adv_fill_template(request, template_id):
    template = get_object_or_404(LegalTemplate, template_id=template_id)

    if request.method == "POST":
        form = DynamicTemplateForm(template.template_content, request.POST)

        if form.is_valid():

            # ✔ Download only the uploaded file
            if template.uploaded_file:
                file_path = template.uploaded_file.path
                file_name = template.uploaded_file.name.split('/')[-1]

                # Read and return the file as download
                with open(file_path, 'rb') as f:
                    response = HttpResponse(
                        f.read(),
                        content_type="application/octet-stream"
                    )
                    response['Content-Disposition'] = f'attachment; filename="%s"' % file_name
                    return response

            else:
                return HttpResponse("No file uploaded by admin.")

    else:
        form = DynamicTemplateForm(template.template_content)

    return render(request, "adv_filltemplate.html", {
        "form": form,
        "template": template
    })



# def admin_template_detailed_view(request, template_id):
#     template = get_object_or_404(LegalTemplate, template_id=template_id)
#     return render(request, "admin_template_detailed_view.html", {"template": template})


def admin_template_detailed_view(request, template_id):
    
        template = LegalTemplate.objects.get(template_id=template_id)
        return render(request, 'admin/admin_template_detailed_view.html', {'template': template})


def template_delete(request, template_id):
    template = LegalTemplate.objects.get(template_id=template_id)
    template.delete()

    messages.success(request, "🗑️ Template deleted successfully!")
    return redirect('template_list')
    




def template_update(request, template_id):

    # Fetch the template by ID
    template = LegalTemplate.objects.get(template_id=template_id)

    if request.method == "POST" and 'template_update' in request.POST:

        template.template_name = request.POST.get('template_name')
        template.category = request.POST.get('category')
        template.description = request.POST.get('description')
        template.template_content = request.POST.get('template_content')

        # ✅ Handle file upload (optional)
        uploaded_file = request.FILES.get('uploaded_file')

        if uploaded_file:
            template.uploaded_file = uploaded_file  # Replace old file

        template.save()

        # Popup and redirect
        messages.success(request, "✅ Template updated successfully!")
        return redirect('admin_template_detailed_view', template_id=template_id)

    return render(request, 'admin/admin_template_detailed_view.html', {'template': template})


from django.http import HttpResponse
from django.shortcuts import render
from .models import legal_office


def add_legal_office(request):
    if request.method == 'POST':

        office_name = request.POST.get('office_name')
        address = request.POST.get('address')
        location = request.POST.get('location')
        contact_no = request.POST.get('contact_no')
        email = request.POST.get('email')
        offences_handled = request.POST.get('offences_handled')
        working_hours = request.POST.get('working_hours')

        # Duplicate check — office email must be unique
        if legal_office.objects.filter(email=email).exists():
            messages.warning(request, "⚠️ Office already exists. Please update.")
            return redirect('add_legal_office')

        # Create new office entry
        legal_office.objects.create(
            office_name=office_name,
            address=address,
            location=location,
            contact_no=contact_no,
            email=email,
            offences_handled=offences_handled,
            working_hours=working_hours
        )

        messages.success(request, "✅ Legal office added successfully!")
        return redirect('add_legal_office')

    return render(request, 'admin/add_legal_office.html')


def admin_view_legal_office(request):
    offices = legal_office.objects.all()
    return render(request, 'admin/admin_view_legal_office.html', {'offices': offices})

def admin_legal_office_detailed_view(request, office_id):
    office = legal_office.objects.get(office_id=office_id)
    return render(request, 'admin/admin_legal_office_detailed_view.html', {'office': office})


def legal_office_delete(request, office_id):
    office = legal_office.objects.get(office_id=office_id)
    office.delete()

    messages.success(request, "🗑️ Legal office deleted successfully!")
    return redirect('admin_view_legal_office')


def legal_office_update(request, office_id):

    # Get office using correct PK
    office = legal_office.objects.get(office_id=office_id)

    if request.method == "POST" and 'legal_office_update' in request.POST:

        office.office_name = request.POST.get('office_name')
        office.address = request.POST.get('address')
        office.location = request.POST.get('location')
        office.contact_no = request.POST.get('contact_no')
        office.email = request.POST.get('email')
        office.offences_handled = request.POST.get('offences_handled')
        office.working_hours = request.POST.get('working_hours')

        office.save()

        # JS popup + redirect
        messages.success(request, "✅ Legal office updated successfully!")
        return redirect('admin_legal_office_detailed_view', office_id=office_id)

    return render(request, 'admin/admin_legal_office_detailed_view.html', {'office': office})


def legal_office_list(request):
    query = request.GET.get('q')

    if query:
        offices = legal_office.objects.filter(
            Q(office_name__icontains=query) |
            Q(address__icontains=query) |
            Q(location__icontains=query) |
            Q(offences_handled__icontains=query)
        )
    else:
        offices = legal_office.objects.all()

    return render(request, 'office_list.html', {
        'offices': offices,
        'query': query
    })

def legal_office_detail(request, pk):
    office = get_object_or_404(legal_office, office_id=pk)
    return render(request, 'office_detail.html', {'office': office})

def adv_legal_office_list(request):
    query = request.GET.get('q')

    if query:
        offices = legal_office.objects.filter(
            Q(office_name__icontains=query) |
            Q(address__icontains=query) |
            Q(location__icontains=query) |
            Q(offences_handled__icontains=query)
        )
    else:
        offices = legal_office.objects.all()

    return render(request, 'adv_office_list.html', {
        'offices': offices,
        'query': query
    })

def adv_legal_office_detail(request, pk):
    office = get_object_or_404(legal_office, office_id=pk)
    return render(request, 'adv_office_detail.html', {'office': office})





def advocates(request):
    specialization = request.GET.get('specialization')
    location = request.GET.get('location')

    advocates_list = Advocate.objects.filter(is_verified=True)

    if specialization:
        advocates_list = advocates_list.filter(specialization__icontains=specialization)
    
    if location:
        advocates_list = advocates_list.filter(location__icontains=location)

    return render(request, "advocates.html", {"advocates": advocates_list})






from django.shortcuts import render, HttpResponse, redirect
from django.contrib.auth.hashers import make_password
from .models import Advocate


def advocate_register(request):
    if request.method == 'POST':
        advocate_name = request.POST.get('advocate_name')
        specialization = request.POST.get('specialization')
        email = request.POST.get('email')
        password = request.POST.get('password')
        experience_years = request.POST.get('experience_years')
        location = request.POST.get('location')
        availability_status = request.POST.get('availability_status')
        state_bar_council= request.POST.get('state_bar_council')
        enrollment_number = request.POST.get('enrollment_no')


        bar_council_idcard= request.FILES.get('bar_council_idcard')
        image = request.FILES.get('image')

        # OPTIONAL: If you want advocates to upload ID proof
        # verification_document = request.FILES.get('verification_document')

        # Get multiple checkbox values
        qualifications_list = request.POST.getlist('qualification')
        qualification = ", ".join(qualifications_list)

        # Convert availability
        availability_status = True if availability_status == "True" else False

        # Check duplicate email
        if Advocate.objects.filter(email=email).exists():
            messages.warning(request, "⚠️ Advocate with this email already exists!")
            return redirect('advocate_register')

        # Hash password
        hashed_password = make_password(password)

        # Save advocate  (➡️ ONLY ADDED is_verified FIELD — NOTHING ELSE CHANGED)
        Advocate.objects.create(
            advocate_name=advocate_name,
            specialization=specialization,
            email=email,
            password=hashed_password,
            experience_years=experience_years,
            location=location,
            availability_status=availability_status,
            image=image,
            qualification=qualification,
            state_bar_council=state_bar_council,
            enrollment_number=enrollment_number,

            bar_council_idcard=bar_council_idcard,

            # ⭐ NEW FIELD ADDED ⭐
            is_verified=False
            # verification_document=verification_document  # optional if you enable it
        )

        # Success popup
        messages.success(request, "✅ Advocate Registered Successfully!")
        return redirect('login')

    return render(request, 'advocate_register.html')








def advocate_login(request):
    if request.method == 'POST':
        email = request.POST.get('email')
        password = request.POST.get('password')

        try:
            advocate = Advocate.objects.get(email=email)
        except Advocate.DoesNotExist:
            messages.error(request, "❌ Email not found!")
            return redirect('advocate_login')

        # Check password
        if not check_password(password, advocate.password):
            messages.error(request, "❌ Incorrect Password!")
            return redirect('advocate_login')

        # ⭐ BLOCK LOGIN IF NOT VERIFIED ⭐
        if not advocate.is_verified:
            messages.warning(request, "⚠️ Your account is not approved by admin yet!")
            return redirect('advocate_login')

        # Login success → store session
        request.session['advocate_name'] = advocate.advocate_name
        request.session['advocate_id'] = advocate.advocate_id


        return redirect('advocate_home')

    return render(request, 'advocate_login.html')



def advocate_profile_view(request):

    # Session variables (set during advocate login)
    advocate_name = request.session.get("advocate_name")
    advocate_id = request.session.get("advocate_id")

    # If not logged in → show popup + redirect
    if not advocate_name or not advocate_id:
        messages.warning(request, "⚠️ Please login to view your profile.")
        return redirect('advocate_login')

    # Load advocate profile data
    advocate = Advocate.objects.get(advocate_id=advocate_id)

    # Load reviews for this advocate
    feedbacks = Feedback.objects.filter(advocate=advocate).order_by('-created_at')

    return render(request, "advocate_profile.html", {
        "advocate": advocate,
        "feedbacks": feedbacks
    })



def pending_advocates(request):
    if not request.session.get("admin_logged_in"):
        return redirect("myadmin")

    advocates = Advocate.objects.filter(is_verified=False)

    return render(request, "admin/pending_advocates.html", {"advocates": advocates})





from django.utils import timezone
from django.template.loader import render_to_string
from django.utils.html import strip_tags
from django.core.mail import send_mail

def approve_advocate(request, advocate_id):
    if not request.session.get("admin_logged_in"):
        return redirect("myadmin")

    advocate = Advocate.objects.get(advocate_id=advocate_id)
    advocate.is_verified = True
    advocate.verified_at = timezone.now()
    advocate.save()

    # send email
    html_message = render_to_string("emails/advocate_approved.html", {
        "advocate": advocate,
        "login_url": request.build_absolute_uri('/login/')
    })
    plain_message = strip_tags(html_message)

    try:
        sent_count = send_mail(
            subject="Your Advocate Account is Approved",
            message=plain_message,
            from_email="no-reply@yourdomain.com",
            recipient_list=[advocate.email],
            html_message=html_message,
            fail_silently=False  # Changed to False to see errors
        )
        print(f"📧 Email sent status: {sent_count}")
        if sent_count == 1:
            messages.success(request, f"✅ Advocate {advocate.advocate_name} approved & Email sent!")
        else:
            messages.warning(request, f"⚠️ Approved, but email might not have sent.")
            
    except Exception as e:
        print(f"❌ Email Sending Error: {e}")
        messages.error(request, f"Approved, but failed to send email: {e}")

    return redirect("pending_advocates")


def verified_advocates(request):
    if not request.session.get("admin_logged_in"):
        return redirect("myadmin")

    advocates = Advocate.objects.filter(is_verified=True)

    return render(request, "admin/verified_advocates.html", {"advocates": advocates})










def advocate_detail(request, id):
    advocate = get_object_or_404(Advocate, advocate_id=id)

    # Default: user has NO approved appointment
    has_approved_appointment = False

    # Check only if user is logged in
    if request.session.get("user_id"):
        has_approved_appointment = Appointment.objects.filter(
            user_id=request.session["user_id"],
            advocate=advocate,
            status="Scheduled"
        ).exists()

    # Fetch reviews
    feedbacks = Feedback.objects.filter(advocate=advocate).order_by('-created_at')

    return render(
        request,
        "advocate_detail.html",
        {
            "adv": advocate,
            "has_approved_appointment": has_approved_appointment,
            "feedbacks": feedbacks
        }
    )

def admin_advocate_profile(request):
    if request.session.get('admin_logged_in'):
        # 🧾 Fetch all registered users
        advocates = Advocate.objects.all()

        return render(request, 'admin/admin_advocate_profile.html', {
            'admin_username': request.session.get('admin_username'),
            'admin_id': request.session.get('admin_user_id'),
            'advocates': advocates,                # 👈 Pass users list to template
        })
    else:
        return HttpResponse("Unauthorized", status=403)
    






def challan_form(request):
    return render(request, "challan_form.html")

def create_challan_payment(request):
    if request.method == "POST":
        challan_no = request.POST.get("challan_no")
        amount = float(request.POST.get("amount"))  # user enters challan fine amount
        
        # Convert to paise for Razorpay
        razor_amount = int(amount * 100)

        # Store details in session for success handler
        request.session["challan_no"] = challan_no
        request.session["amount"] = razor_amount

        # Create Razorpay Order
        currency = "INR"
        razorpay_order = razorpay_client.order.create(dict(
            amount = razor_amount,
            currency = currency,
            payment_capture = '0'
        ))

        razorpay_order_id = razorpay_order["id"]
        callback_url = "/challan/paymenthandler/"

        context = {
            "razorpay_order_id": razorpay_order_id,
            "razorpay_merchant_key": settings.RAZORPAY_KEY_ID,
            "razorpay_amount": razor_amount,
            "display_amount": amount,
            "currency": currency,
            "callback_url": callback_url,
            "challan_no": challan_no,
        }

        return render(request, "challan_payment.html", context)

    return render(request, "challan_form.html")


@csrf_exempt
def challan_paymenthandler(request):

    if request.method == "POST":

        payment_id = request.POST.get("razorpay_payment_id")
        order_id = request.POST.get("razorpay_order_id")
        signature = request.POST.get("razorpay_signature")

        amount = request.session.get("amount")
        challan_no = request.session.get("challan_no")

        # Get user
        user_email = request.session.get("user_email")
        user = UserData.objects.get(email=user_email)

        params = {
            "razorpay_order_id": order_id,
            "razorpay_payment_id": payment_id,
            "razorpay_signature": signature
        }

        try:
            # STEP 1: Verify signature
            razorpay_client.utility.verify_payment_signature(params)

            # STEP 2: Check existing payment status
            payment_info = razorpay_client.payment.fetch(payment_id)

            if payment_info["status"] != "captured":
                # Capture only if NOT captured
                razorpay_client.payment.capture(payment_id, amount)

            # STEP 3: Store success in DB
            Payment.objects.create(
                user_id=user,
                challan_no=challan_no,
                amount=amount / 100,
                status=True
            )

            # Remove temporary session data
            request.session.pop("amount", None)
            request.session.pop("challan_no", None)

            return redirect("challan_success")

        except Exception as e:
            print("PAYMENT ERROR:", e)

            # Store failed payment attempt
            Payment.objects.create(
                user_id=user,
                challan_no=challan_no,
                amount=amount / 100,
                status=False
            )

            return redirect("challan_failed")

    return HttpResponseBadRequest()




def challan_success(request):
    return render(request, 'challan_success.html')


def challan_failed(request):
    return render(request, 'challan_failed.html')


def admin_challan_detail(request, transaction_id):
    payment = Payment.objects.get(transaction_id=transaction_id)
    return render(request, 'admin_challan_detail.html', {'payment': payment})








    





# -----------------------------
# ADMIN CHALLAN LIST VIEW
# -----------------------------

def admin_challan_view(request):
    # Search value
    search = request.GET.get("search", "")

    # Always show only paid payments
    payments = Payment.objects.filter(status=True).order_by("-payment_date")

    # Search by challan number or user name
    if search:
        payments = payments.filter(
            challan_no__icontains=search
        ) | payments.filter(
            user_id__fullname__icontains=search
        )

    context = {
        "payments": payments,
        "search": search,
    }
    return render(request, "admin/admin_challan_view.html", context)



# -----------------------------
# ADMIN CHALLAN DETAIL VIEW
# -----------------------------

# def admin_challan_detail(request, tx_id):
#     payment = get_object_or_404(Payment, transaction_id=tx_id)

#     context = {
#         "payment": payment
#     }
#     return render(request, "admin_challan_detail.html", context)







from django.shortcuts import render, get_object_or_404
from django.db.models import Q
from .models import Law
from django.http import HttpResponse, JsonResponse

def law_list(request):
    query = request.GET.get('q')
    if query:
        laws = Law.objects.filter(
            Q(title__icontains=query) | Q(description__icontains=query) | Q(category__icontains=query)
        )
    else:
        laws = Law.objects.all()
    
    return render(request, 'laws/law_list.html', {'laws': laws, 'query': query})

def law_detail(request, pk):
    law = get_object_or_404(Law, pk=pk)
    return render(request, 'laws/law_detail.html', {'law': law})

def download_law(request, pk):
    law = get_object_or_404(Law, pk=pk)
    content = f"""Title: {law.title}
Category: {law.category}

Description:
{law.description}

---
Penalty: {law.punishment}
Status: {law.bailable_status}, {law.cognizable_status}
Court: {law.triable_by}
"""
    response = HttpResponse(content, content_type='text/plain')
    response['Content-Disposition'] = f'attachment; filename="{law.title}.txt"'
    return response

from django.contrib.auth.decorators import login_required

@login_required(login_url='login')
def chat_view(request):
    return render(request, 'laws/chat.html')


import os
import google.generativeai as genai
from dotenv import load_dotenv

load_dotenv()

# Configure Gemini
GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')
if GOOGLE_API_KEY:
    genai.configure(api_key=GOOGLE_API_KEY)


from django.http import JsonResponse
from .models import Law
import re
import google.generativeai as genai


# (Duplicate chatbot removed to use the enhanced version below)



from datetime import datetime, date
from django.http import HttpResponse
from django.shortcuts import get_object_or_404, redirect
from .models import Advocate, Appointment, UserData


def book_appointment(request, advocate_id):
    advocate = get_object_or_404(Advocate, advocate_id=advocate_id)

    # ----------------------------------------------------
    # ✅ CHECK IF USER IS LOGGED IN (SESSION BASED)
    # ----------------------------------------------------
    if not request.session.get("user_id"):
        return HttpResponse("""
        <html>
        <body>
            <script>
                window.onload = function() {
                    const msg = document.createElement('div');
                    msg.textContent = "⚠️ You need to login first to book an appointment!";
                    msg.style.position = 'fixed';
                    msg.style.top = '50%';
                    msg.style.left = '50%';
                    msg.style.transform = 'translate(-50%, -50%)';
                    msg.style.background = '#D9534F';
                    msg.style.color = 'white';
                    msg.style.padding = '25px 50px';
                    msg.style.borderRadius = '10px';
                    msg.style.fontSize = '22px';
                    msg.style.zIndex = '9999';
                    msg.style.boxShadow = '0 4px 10px rgba(0,0,0,0.3)';
                    msg.style.fontFamily = 'Poppins, sans-serif';
                    msg.style.opacity = '0';
                    msg.style.transition = 'opacity 0.5s ease';
                    document.body.appendChild(msg);

                    setTimeout(() => { msg.style.opacity = '1'; }, 100);

                    setTimeout(() => {
                        msg.style.opacity = '0';
                        setTimeout(() => {
                            msg.remove();
                            window.location.href = '/login/';
                        }, 400);
                    }, 1500);
                }
            </script>
        </body>
        </html>
        """)

    # ----------------------------------------------------
    # ✅ CREATE APPOINTMENT WITH VALIDATION
    # ----------------------------------------------------
    if request.method == "POST":
        appointment_date_str = request.POST.get("appointment_date")

        # ❌ Date not selected
        if not appointment_date_str:
            return HttpResponse("<script>alert('Please select an appointment date');history.back();</script>")

        # ❌ Invalid date format
        try:
            appointment_date = datetime.strptime(appointment_date_str, "%Y-%m-%d").date()
        except ValueError:
            return HttpResponse("<script>alert('Invalid date format');history.back();</script>")

        # ❌ Past date
        if appointment_date < date.today():
            return HttpResponse("<script>alert('You cannot book an appointment in the past');history.back();</script>")

        user = get_object_or_404(UserData, id=request.session["user_id"])

        # ❌ Duplicate appointment check
        if Appointment.objects.filter(
            user=user,
            advocate=advocate,
            appointment_date=appointment_date
        ).exists():
            return HttpResponse("<script>alert('You already have an appointment on this date');history.back();</script>")

        # ✅ Create appointment
        Appointment.objects.create(
            user=user,
            advocate=advocate,
            appointment_date=appointment_date,
            status="Pending"
        )

        # ----------------------------------------------------
        # ✅ SUCCESS MESSAGE
        # ----------------------------------------------------
        return HttpResponse(f"""
        <html>
        <body>
        <script>
            window.onload = function() {{
                const msg = document.createElement('div');
                msg.textContent = "✅ Appointment request sent successfully!";
                msg.style.position = 'fixed';
                msg.style.top = '50%';
                msg.style.left = '50%';
                msg.style.transform = 'translate(-50%, -50%)';
                msg.style.background = '#C89D66';
                msg.style.color = 'white';
                msg.style.padding = '25px 50px';
                msg.style.borderRadius = '10px';
                msg.style.fontSize = '22px';
                msg.style.zIndex = '9999';
                msg.style.boxShadow = '0 4px 10px rgba(0,0,0,0.3)';
                msg.style.fontFamily = 'Poppins, sans-serif';
                msg.style.opacity = '0';
                msg.style.transition = 'opacity 0.5s ease';
                document.body.appendChild(msg);

                setTimeout(() => {{ msg.style.opacity = '1'; }}, 100);

                setTimeout(() => {{
                    msg.style.opacity = '0';
                    setTimeout(() => {{
                        msg.remove();
                        window.location.href = '/advocate/{advocate.advocate_id}/';
                    }}, 400);
                }}, 1500);
            }};
        </script>
        </body>
        </html>
        """)

    return redirect("advocate_detail", advocate_id=advocate.advocate_id)



def advocate_logout(request):
    request.session.flush()
    return redirect('login')




def advocate_appointments(request):
    if not request.session.get("advocate_id"):
        return redirect("advocate_login")

    advocate = get_object_or_404(
        Advocate,
        advocate_id=request.session["advocate_id"]
    )

    appointments = Appointment.objects.filter(
        advocate=advocate
    ).order_by("-appointment_date")

    return render(request, "appointments.html", {
        "appointments": appointments
    })


def approve_appointment(request, appointment_id):
    if not request.session.get("advocate_id"):
        return redirect("advocate_login")

    appointment = get_object_or_404(Appointment, appointment_id=appointment_id)

    if appointment.advocate.advocate_id != request.session["advocate_id"]:
        return redirect("advocate_appointments")

    if appointment.status != "Pending":
        return redirect("advocate_appointments")

    appointment.status = "Scheduled"
    appointment.save()

    return redirect("advocate_appointments")


def reject_appointment(request, appointment_id):
    if not request.session.get("advocate_id"):
        return redirect("advocate_login")

    appointment = get_object_or_404(Appointment, appointment_id=appointment_id)

    if appointment.advocate.advocate_id != request.session["advocate_id"]:
        return redirect("advocate_appointments")

    if appointment.status != "Pending":
        return redirect("advocate_appointments")

    appointment.status = "Cancelled"
    appointment.save()

    return redirect("advocate_appointments")



def user_appointments(request):
    if not request.session.get("user_id"):
        return redirect("login")

    user = UserData.objects.get(id=request.session["user_id"])

    appointments = Appointment.objects.filter(
        user=user
    ).order_by("-appointment_date")

    return render(
        request,
        "user_appointments.html",
        {"appointments": appointments}
    )



from django.shortcuts import render
from django.http import HttpResponse
from .models import Awareness_content
from .utils import extract_youtube_video_id, get_youtube_embed_data











from urllib.parse import urlparse, parse_qs


def extract_video_id(url):
    try:
        if "youtu.be" in url:
            return url.split("/")[-1].split("?")[0]

        parsed = urlparse(url)

        if "youtube.com" in parsed.netloc:
            if parsed.path.startswith("/shorts/"):
                return parsed.path.split("/shorts/")[1].split("/")[0]

            if parsed.path == "/watch":
                return parse_qs(parsed.query).get("v", [None])[0]
    except Exception:
        return None

    return None


def add_awareness_content(request):
    if request.method == 'POST':
        title = request.POST.get('title')
        link = request.POST.get('link')
        category = request.POST.get('category')

        # 🔒 Force content type
        content_type = "YouTube Shorts"

        # 🎯 Validate & normalize link
        video_id = extract_youtube_video_id(link)


        if not video_id:
            return HttpResponse("""
<html><body>
<script>
alert("Invalid YouTube Shorts link");
window.history.back();
</script>
</body></html>
""")

        # 🔁 Prevent duplicate videos (even if link format differs)
        if Awareness_content.objects.filter(link__icontains=video_id).exists():
            return HttpResponse("""
<html>
<body>
<script>
window.onload = function() {
    const msg = document.createElement('div');
    msg.textContent = "⚠️ Content already exists. Please update.";
    msg.style.position = 'fixed';
    msg.style.top = '50%';
    msg.style.left = '50%';
    msg.style.transform = 'translate(-50%, -50%)';
    msg.style.background = '#C89D66';
    msg.style.color = 'white';
    msg.style.padding = '25px 50px';
    msg.style.borderRadius = '10px';
    msg.style.fontSize = '22px';
    msg.style.zIndex = '9999';
    msg.style.boxShadow = '0 4px 10px rgba(0,0,0,0.3)';
    msg.style.fontFamily = 'Poppins, sans-serif';
    msg.style.opacity = '0';
    msg.style.transition = 'opacity 0.5s ease';
    document.body.appendChild(msg);

    setTimeout(() => { msg.style.opacity = '1'; }, 100);
    setTimeout(() => {
        msg.style.opacity = '0';
        setTimeout(() => {
            msg.remove();
            window.history.back();
        }, 400);
    }, 1500);
}
</script>
</body>
</html>
""")

        # ✅ Save content
        Awareness_content.objects.create(
            title=title,
            type=content_type,
            link=link,
            category=category
        )

        return HttpResponse("""
<html>
<body>
<script>
window.onload = function() {
    const msg = document.createElement('div');
    msg.textContent = "✅ Awareness content added successfully!";
    msg.style.position = 'fixed';
    msg.style.top = '50%';
    msg.style.left = '50%';
    msg.style.transform = 'translate(-50%, -50%)';
    msg.style.background = '#C89D66';
    msg.style.color = 'white';
    msg.style.padding = '25px 50px';
    msg.style.borderRadius = '10px';
    msg.style.fontSize = '22px';
    msg.style.zIndex = '9999';
    msg.style.boxShadow = '0 4px 10px rgba(0,0,0,0.3)';
    msg.style.fontFamily = 'Poppins, sans-serif';
    msg.style.opacity = '0';
    msg.style.transition = 'opacity 0.5s ease';
    document.body.appendChild(msg);

    setTimeout(() => { msg.style.opacity = '1'; }, 100);
    setTimeout(() => {
        msg.style.opacity = '0';
        setTimeout(() => {
            msg.remove();
            window.location.href = '/add-awareness/';
        }, 400);
    }, 1500);
}
</script>
</body>
</html>
""")

    return render(request, 'admin/add_awareness.html')

from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required

from .models import LegalDocument, UserData
from .utils import (
    extract_text_from_document,
    analyze_legal_document_with_gemini
)


@login_required(login_url='login')
def legal_document_analysis(request):
    ai_summary = None
    error = None
    loading = False

    # ✅ Get UserData from session
    user_id = request.session.get('user_id')
    if not user_id:
        return redirect('login')

    try:
        user = UserData.objects.get(id=user_id)
    except UserData.DoesNotExist:
        return redirect('login')

    if request.method == 'POST':
        uploaded_file = request.FILES.get('document')

        if not uploaded_file:
            error = "Please upload a document."
        else:
            loading = True

            # ✅ Save document with valid user
            document = LegalDocument.objects.create(
                user=user,
                file=uploaded_file
            )

            try:
                extracted_text = extract_text_from_document(
                    document.file.path
                )

                if not extracted_text.strip():
                    error = "Could not extract readable text from the document."

                else:
                    # ✅ Gemini-powered analysis
                    ai_summary = analyze_legal_document_with_gemini(
                        extracted_text
                    )

                    if ai_summary:
                        document.ai_summary = ai_summary
                        document.save()
                    else:
                        error = "AI service is temporarily busy (Rate Limit reached). Please wait a few minutes and try again."

            except Exception as e:
                print("Document Analysis Error:", e)
                error = "Error processing the document. Please try again."

    return render(request, 'legal_document_analysis.html', {
        'ai_summary': ai_summary,
        'error': error,
        'loading': False  # ✅ DO NOT SHOW LOADING BOX ONCE RENDERED
    })



from django.http import HttpResponse
from django.contrib.auth.decorators import login_required
from django.views.decorators.http import require_POST
from django.core.mail import send_mail
from django.conf import settings

from .models import UserData


@login_required
@require_POST
def send_sos(request):
    user = request.user

    try:
        userdata = UserData.objects.get(email=user.email)
    except UserData.DoesNotExist:
        return HttpResponse("""
<html><body>
<script>
alert("User details not found");
window.history.back();
</script>
</body></html>
""")

    subject = "🚨 EMERGENCY ALERT – LAWLY"

    message = f"""
🚨 EMERGENCY ALERT 🚨

Triggered By:
--------------------
Name     : {userdata.fullname}
Email    : {userdata.email}
Phone    : {userdata.contact_no}
Location : {userdata.location}

Please respond immediately.
"""

    
    # Save to Database
    SOS.objects.create(
        user=userdata,
        location=userdata.location,
        message=message
    )

    send_mail(
        subject=subject,
        message=message,
        from_email=settings.DEFAULT_FROM_EMAIL,
        recipient_list=["sreenath0880@gmail.com"],
        fail_silently=False,
    )

    # ✅ Custom popup with OK button (NO auto-close)
    return HttpResponse("""
<html>
<body>
<script>
window.onload = function () {
    const overlay = document.createElement('div');
    overlay.style.position = 'fixed';
    overlay.style.inset = '0';
    overlay.style.background = 'rgba(0,0,0,0.6)';
    overlay.style.display = 'flex';
    overlay.style.alignItems = 'center';
    overlay.style.justifyContent = 'center';
    overlay.style.zIndex = '9999';

    const box = document.createElement('div');
    box.style.background = '#fff';
    box.style.padding = '30px';
    box.style.borderRadius = '12px';
    box.style.boxShadow = '0 10px 25px rgba(0,0,0,0.2)';
    box.style.textAlign = 'center';
    box.style.maxWidth = '400px';
    box.style.width = '90%';
    box.style.fontFamily = "'Poppins', sans-serif";

    const icon = document.createElement('div');
    icon.innerHTML = '🚨';
    icon.style.fontSize = '50px';
    icon.style.marginBottom = '15px';

    const title = document.createElement('h2');
    title.innerText = 'SOS Sent Successfully!';
    title.style.margin = '0 0 10px 0';
    title.style.color = '#333';
    title.style.fontSize = '22px';

    const msg = document.createElement('p');
    msg.innerText = 'Your emergency alert has been sent to the admin and support team.';
    msg.style.color = '#666';
    msg.style.marginBottom = '20px';
    msg.style.fontSize = '15px';
    msg.style.lineHeight = '1.5';

    const btn = document.createElement('button');
    btn.innerText = 'OK';
    btn.style.background = '#d9534f';
    btn.style.color = '#fff';
    btn.style.border = 'none';
    btn.style.padding = '10px 30px';
    btn.style.borderRadius = '6px';
    btn.style.fontSize = '16px';
    btn.style.cursor = 'pointer';
    btn.style.fontWeight = '600';
    
    btn.onmouseover = function() { btn.style.background = '#c9302c'; };
    btn.onmouseout = function() { btn.style.background = '#d9534f'; };

    btn.onclick = function () {
        document.body.removeChild(overlay);
        window.history.back();
    };

    box.appendChild(icon);
    box.appendChild(title);
    box.appendChild(msg);
    box.appendChild(btn);
    overlay.appendChild(box);
    document.body.appendChild(overlay);
};
</script>
</body>
</html>
""")

def admin_sos_alerts(request):
    alerts = SOS.objects.all().order_by('-created_at')
    return render(request, "admin/admin_sos_alerts.html", {"alerts": alerts})










def delete_advocate(request, advocate_id):
    if not request.session.get("admin_logged_in"):
        return redirect("myadmin")

    advocate = get_object_or_404(Advocate, advocate_id=advocate_id)
    advocate.delete()

    messages.success(request, f"🗑️ Advocate {advocate.advocate_name} has been deleted successfully!")
    
    # Redirect back to the page request came from (Pending or Verified list)
    return redirect(request.META.get('HTTP_REFERER', 'verified_advocates'))

def chat_api(request):
    if request.method != 'GET':
        return JsonResponse({'response': 'Invalid request'}, status=400)

    query = request.GET.get('message', '').strip()
    if not query:
        return JsonResponse({'response': 'Please ask me about a law.'})

    text = query.lower()

    # Casual Greetings
    casual_patterns = [
        r'\b(hi|hello|hey|good morning|good evening|good afternoon)\b',
        r'\b(thank|thanks|thx|appreciate)\b',
        r'\b(how are you|what\'s up)\b',
        r'(👏|😊|🙂|😀|👍)'
    ]
    if any(re.search(p, text, re.I) for p in casual_patterns):
        if any(w in text for w in ['hi', 'hello', 'hey', 'ഹായ്', 'ഹലോ']):
            return JsonResponse({'response': "Hello! I am Lawly, your Legal Assistant. How can I help you today?"})
        return JsonResponse({'response': "You're welcome! Feel free to ask more legal questions."})

    # Forced BNS Mode
    section_pattern = r'\b(bns\s*section\s*\d+|section\s*\d+|bns\s*\d+)\b'
    force_bns_answer = bool(re.search(section_pattern, text))

    # Keywords and Stemming
    query_words = text.split()
    stop_words = {'i','me','my','myself','we','our','what','which','who','whom','this','that','is','are','was','were','be','been','being','have','has','had','do','does','did','a','an','the','and','but','if','or','because','as','until','while','of','at','by','for','with','about','against','between','into','through','during','before','after','above','below','to','from','up','down','in','out','on','off','over','under','again','further','then','once','here','there','when','where','why','how','all','any','both','each','few','more','most','other','some','such','no','nor','not','only','own','same','so','than','too','very','can','will','just','don','should','now','please','tell','me'}
    keywords = [w for w in query_words if w not in stop_words]
    stems = {kw for kw in keywords}
    for kw in keywords:
        if len(kw) > 4:
            if kw.endswith('ing'): stems.add(kw[:-3])
            elif kw.endswith('ed'): stems.add(kw[:-2])
            elif kw.endswith('s'): stems.add(kw[:-1])
            elif kw.endswith('es'): stems.add(kw[:-2])
            elif kw.endswith('ment'): stems.add(kw[:-4])
    
    # Keyword Expansion (Add common synonyms/related terms)
    expansion_map = {
        'accident': ['motor', 'traffic', 'vehicle', 'driving', 'collision', 'road'],
        'theft': ['stealing', 'robbery', 'burglary', 'property'],
        'murder': ['homicide', 'killing', 'death', 'body'],
        'cyber': ['internet', 'online', 'computer', 'fraud'],
        'family': ['marriage', 'divorce', 'domestic', 'children'],
        'money': ['financial', 'fraud', 'bribe', 'corruption'],
    }
    for kw in list(stems):
        if kw in expansion_map:
            stems.update(expansion_map[kw])

    from django.db.models import Q
    scored_laws = []
    if stems:
        search_filter = Q()
        for s in stems:
            if len(s) < 3: continue
            search_filter |= Q(title__icontains=s) | Q(description__icontains=s) | Q(category__icontains=s)
        db_laws = Law.objects.filter(search_filter).distinct()
        for law in db_laws:
            score = 0
            l_title, l_desc = law.title.lower(), law.description.lower()
            if text in l_title: score += 50
            if text in l_desc: score += 10
            for s in stems:
                if len(s) < 3: continue
                if s in l_title: score += 20
                if s in l_desc: score += 10
                if s in law.category.lower(): score += 10
            scored_laws.append((score, law))
    scored_laws.sort(key=lambda x: x[0], reverse=True)

    # Gemini Integration
    from django.conf import settings
    import google.generativeai as genai
    from .models import AIChatCache
    
    # Check Cache
    cached = AIChatCache.objects.filter(query=query).first()
    if cached:
        print("Returning cached AI response.")
        return JsonResponse({'response': cached.response})

    api_keys = getattr(settings, 'GEMINI_API_KEYS', [settings.GEMINI_API_KEY] if hasattr(settings, 'GEMINI_API_KEY') and settings.GEMINI_API_KEY else [])
    
    if api_keys:
        if scored_laws and not force_bns_answer:
            top_laws = scored_laws[:3]
            context_text = "\n".join([f"Title: {l.title}\nDescription: {l.description}\nPunishment: {l.punishment}\nBailable: {l.bailable_status}" for s, l in top_laws])
            prompt = f"Expert Legal Assistant. Context:\n{context_text}\nQuestion: {query}\nInstructions: Explain simply using HTML tags (p, strong, ul, li, em)."
        else:
            prompt = f"Expert Legal Assistant. Question: {query}. The database did not match. Answer from your own knowledge using BNS preferred. Use HTML tags."

        import time
        max_retries = 2
        for attempt in range(max_retries):
            for idx, key in enumerate(api_keys):
                if not key: continue
                try:
                    genai.configure(api_key=key)
                    # Switched to 2.0-flash for availability/compatibility
                    try:
                        model = genai.GenerativeModel('gemini-2.0-flash')
                        response = model.generate_content(prompt)
                    except Exception as model_err:
                        print(f"Model gemini-2.0-flash failed in Chat: {model_err}. Trying gemini-flash-latest...")
                        model = genai.GenerativeModel('gemini-flash-latest')
                        response = model.generate_content(prompt)
                    
                    if response and response.text:
                        # Save to Cache
                        AIChatCache.objects.create(query=query, response=response.text)
                        return JsonResponse({'response': response.text})
                except Exception as e:
                    print(f"Gemini Chat Error (Key {idx+1}, Attempt {attempt+1}):", e)
                    if "429" in str(e):
                        if idx < len(api_keys) - 1:
                            print(f"Key {idx+1} rate limited. Retrying with key {idx+2}...")
                            continue
                        elif attempt < max_retries - 1:
                            print("All keys rate limited. Waiting 10 seconds to retry queue...")
                            time.sleep(10)  # Sequential Queue wait
                            break  # Exit key loop to retry from Key 1
                        else:
                            return JsonResponse({'response': "<p><strong>AI service is temporarily busy (All keys hit Rate Limit). Please wait 60 seconds and try again.</strong></p>"})
                    
                    # For other errors, try the next key if available
                    if idx < len(api_keys) - 1:
                        print(f"Key {idx+1} failed with {type(e).__name__}. Trying key {idx+2}...")
                        continue
                    
                    # If it's the last attempt and last key, then simply break to local fallback
                    if attempt == max_retries - 1:
                        break

    # Fallback HTML
    if scored_laws and not force_bns_answer:
        response_text = 'I found some relevant laws matching your query:<br><ul>'
        for score, law in scored_laws[:5]:
            punishment_text = (f'<br><div style=margin-top:8px; padding:8px; background:rgba(239,68,68,0.1); border-left:3px solid #ef4444; border-radius:4px;><strong style=color:#ef4444;>Penalty:</strong> {law.punishment}</div>' if law.punishment else '')
            response_text += f"""<li style=margin-bottom: 24px;><div style=display:flex; justify-content:space-between;><strong style=color:var(--accent-color); font-size:1.1rem;>{law.title}</strong><a href=/law/{law.pk}/download/ target=_blank style=font-size:0.8rem; background:var(--primary-color); color:white; padding:4px 12px; border-radius:6px; text-decoration:none;> Download</a></div><div>{law.description[:400]}...</div>{punishment_text}</li>"""
        response_text += '</ul>'
        return JsonResponse({'response': response_text})

    # Default Result
    categories = ', '.join(set(Law.objects.values_list('category', flat=True).distinct()))
    return JsonResponse({'response': f"No laws matched '{query}'.<br>My database includes: <b>{categories}</b>.<br>Try keywords like theft, cybercrime, or murder."})


@csrf_exempt
def download_chat(request):
    if request.method == 'POST':
        content = request.POST.get('content', '')
        
        # Clean HTML tags
        clean_text = strip_tags(content)
        
        # Create Document
        doc = Document()
        doc.add_heading('LawBot Legal Advice', 0)
        
        # Add content
        doc.add_paragraph(clean_text)
        
        # Add Footer Disclaimer
        doc.add_paragraph('\n--------------------------------------------------')
        doc.add_paragraph('Disclaimer: This document is AI-generated and does not constitute professional legal advice. Please consult a licensed advocate for official legal counsel.')
        
        # Prepare Response
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="Legal_Advice.docx"'
        
        doc.save(response)
        return response
    
    return HttpResponseBadRequest("Invalid request")
